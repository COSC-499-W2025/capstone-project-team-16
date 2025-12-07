# Builds both a plain text summary and a Word resume

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt


def _build_project_line(p: dict) -> str:
    """
    building a short resmume desc about a project.
    (same idea as in alternative_analysis, but kept local so this
    file can be imported independently.)
    """
    name = p["project"]
    langs = p.get("languages", "Unknown")
    frameworks = p.get("frameworks", "None")
    duration = p.get("duration_days", 0)
    code_files = p.get("code_files", 0)
    test_files = p.get("test_files", 0)
    project_type = p.get("project_type", "software")

    pieces = []

    main = f"Contributed to project '{name}'"
    if project_type and project_type != "Unknown":
        main = f"Contributed to {project_type.lower()} project '{name}'"
    if langs and langs != "Unknown":
        main += f" using {langs}"
    pieces.append(main)

    details = []
    if code_files:
        details.append(f"{code_files} code files")
    if test_files:
        details.append(f"{test_files} test files")
    if duration:
        details.append(f"over {duration} days")

    if details:
        pieces.append("worked on " + ", ".join(details))

    if frameworks and frameworks != "None":
        pieces.append(f"with frameworks such as {frameworks}")

    return "; ".join(pieces) + "."


def _write_txt_summary(
    txt_path: str,
    top_projects: list[dict],
    chronological_projects: list[dict],
    skills_output: list[dict],
) -> None:
    """
    plain-text summary for debugging / quick copy paste. we can remove this if we don't need a text file.
    """
    lines: list[str] = []
    lines.append("PROJECT PORTFOLIO SUMMARY")
    lines.append("=" * 60)
    lines.append("")

    # Top projects
    lines.append("Top Projects")
    lines.append("-" * 40)
    for p in top_projects:
        lines.append(f"- {_build_project_line(p)}")
    lines.append("")

    # Project timeline
    lines.append("Chronological Project Timeline")
    lines.append("-" * 40)
    for p in chronological_projects:
        lines.append(
            f"- {p['name']} – {p['first_used']} → {p['last_used']}"
        )
    lines.append("")

    # Skills over time
    lines.append("Skills Used Over Time")
    lines.append("-" * 40)
    for row in skills_output:
        lines.append(
            f"- {row['skill']} – {row['first_used']} → {row['last_used']}"
        )
    lines.append("")

    os.makedirs(os.path.dirname(txt_path), exist_ok=True)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def _write_docx_resume(
    docx_path: str,
    top_projects: list[dict],
    chronological_projects: list[dict],
    skills_output: list[dict],
) -> None:

    """
    Word document
    """
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)

    doc = Document()

    # --- Header / title ---
    title = doc.add_heading("Project Portfolio Resume", level=0)
    title.runs[0].font.size = Pt(20)

    # Placeholder for user info (later on, when we store user infos this need to be updated)
    info_p = doc.add_paragraph()
    info_p.add_run("Name: ").bold = True
    info_p.add_run("  |  Email: .gmail.com  |  GitHub: github.com/")
    info_p.style.font.size = Pt(10)

    doc.add_paragraph()  # blank line

    # --- Top Projects section ---
    doc.add_heading("Top Projects", level=1)

    if not top_projects:
        doc.add_paragraph("No projects detected.", style="List Bullet")
    else:
        for p in top_projects:
            # First line: bold project name + timeframe
            first_date = (
                p.get("first_modified").date().isoformat()
                if isinstance(p.get("first_modified"), datetime)
                else ""
            )
            last_date = (
                p.get("last_modified").date().isoformat()
                if isinstance(p.get("last_modified"), datetime)
                else ""
            )

            para = doc.add_paragraph(style="List Bullet")
            run_name = para.add_run(p["project"])
            run_name.bold = True

            if first_date and last_date:
                para.add_run(f"  ({first_date} – {last_date})")

            # Second line (same bullet) – description
            para.add_run("\n" + _build_project_line(p))

    doc.add_paragraph()  # spacing

    # --- Chronological Project Timeline ---
    doc.add_heading("Project Timeline", level=1)
    for p in chronological_projects:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(p["name"]).bold = True
        para.add_run(f" – {p['first_used']} → {p['last_used']}")

    doc.add_paragraph()

    # --- Skills Used Over Time ---
    doc.add_heading("Skills Used Over Time", level=1)

    if skills_output:
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Skill"
        hdr_cells[1].text = "First Used"
        hdr_cells[2].text = "Last Used"

        for row in skills_output:
            row_cells = table.add_row().cells
            row_cells[0].text = row["skill"]
            row_cells[1].text = row["first_used"]
            row_cells[2].text = row["last_used"]
    else:
        doc.add_paragraph("No skills detected.", style="List Bullet")

    # footer note
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot_run = foot.add_run(
        "Generated automatically from code repositories using Skill Scope."
    )
    foot_run.italic = True
    foot_run.font.size = Pt(8)

    doc.save(docx_path)


def generate_resume(
    project_summaries: list[dict],
    chronological_projects: list[dict],
    skills_output: list[dict],
) -> None:
    """
    Main entry point called from alternative_analysis.analyze_projects.
    """
    # Sort projects by score again just in case
    sorted_projects = sorted(
        project_summaries, key=lambda p: p.get("score", 0), reverse=True
    )
    top_projects = sorted_projects[:3]

    from file_parser import OUTPUT_DIR

    txt_path = os.path.join(OUTPUT_DIR, "resume_output.txt")
    docx_path = os.path.join(OUTPUT_DIR, "portfolio_resume.docx")


    _write_txt_summary(txt_path, top_projects, chronological_projects, skills_output)
    _write_docx_resume(docx_path, top_projects, chronological_projects, skills_output)

    print(f"saved résumé text to {txt_path}")
    print(f"saved résumé document to {docx_path}")
